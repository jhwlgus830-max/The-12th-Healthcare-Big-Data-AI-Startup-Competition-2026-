"""
build_knowledge_db.py — RAG 상담 지식 DB 구축 자동화 스크립트
============================================================
rag_documents/ 폴더에 위치한 모든 PDF, DOCX, PNG, MD 파일을 자동으로 감지하여 텍스트를 추출하고,
FAISS 벡터 DB(vector_store/clinical_kb)를 빌드합니다.
또한, '보건복지부 국립정신건강센터_정신건강 관련기관 정보' CSV 파일을 로드하여 
표준 지역명(region) 메타데이터와 함께 public_resource_kb 인덱스로 빌드합니다.
"""

import os
import sys
import csv
import re
from typing import List
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Windows 콘솔 유니코드 출력 인코딩 설정
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

# data 폴더를 path에 추가하여 로컬 모듈 로드 지원
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from rag_engine import RAGEngine

# 기본 경로 정의
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, "rag_documents")
VECTOR_STORE_DIR = os.path.join(BASE_DIR, "vector_store")

_ocr_reader = None

def get_ocr_reader():
    """EasyOCR 리더를 싱글톤 패턴으로 초기화하여 메모리 및 실행 속도를 최적화합니다."""
    global _ocr_reader
    if _ocr_reader is None:
        os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
        import easyocr
        _ocr_reader = easyocr.Reader(['ko', 'en'], gpu=False)
    return _ocr_reader

def extract_text_from_pdf(pdf_path: str) -> str:
    """PDF 파일로부터 전체 텍스트를 정밀하게 추출합니다."""
    print(f"  └─ PDF 텍스트 추출 중: {os.path.basename(pdf_path)}")
    try:
        reader = PdfReader(pdf_path)
        full_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text.append(text)
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"[Error] PDF 파싱 중 오류 발생: {pdf_path}. 상세: {e}")
        return ""

def extract_text_from_docx(docx_path: str) -> str:
    """DOCX 파일로부터 본문 문단 및 표(Table) 내부의 텍스트를 정밀하게 추출합니다."""
    print(f"  └─ DOCX 텍스트 추출 중: {os.path.basename(docx_path)}")
    try:
        doc = DocxDocument(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"[Error] DOCX 파싱 중 오류 발생: {docx_path}. 상세: {e}")
        return ""

def extract_text_from_png(png_path: str) -> str:
    """PNG 이미지 파일로부터 PIL을 사용하여 로딩한 뒤 EasyOCR로 한글/영어 텍스트를 추출합니다."""
    print(f"  └─ 이미지(OCR) 텍스트 추출 중: {os.path.basename(png_path)}")
    try:
        from PIL import Image
        import numpy as np
        reader = get_ocr_reader()
        with Image.open(png_path) as img:
            img = img.convert('RGB')
            img_np = np.array(img)
        result = reader.readtext(img_np)
        text_lines = [line[1].strip() for line in result if line[1].strip()]
        return "\n".join(text_lines)
    except Exception as e:
        print(f"[Error] PNG OCR 중 오류 발생: {png_path}. 상세: {e}")
        return ""

def extract_text_from_md(md_path: str) -> str:
    """MD 파일로부터 전체 텍스트를 정밀하게 추출합니다."""
    print(f"  └─ MD 텍스트 추출 중: {os.path.basename(md_path)}")
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print(f"[Error] MD 파싱 중 오류 발생: {md_path}. 상세: {e}")
        return ""

def process_and_chunk_documents() -> List[Document]:
    """rag_documents 폴더 안의 모든 PDF, DOCX, PNG, MD 파일을 청킹하여 Document 객체 리스트로 반환합니다."""
    if not os.path.exists(DOCS_DIR):
        print(f"[Warning] 문서 폴더가 존재하지 않아 새로 생성합니다: {DOCS_DIR}")
        os.makedirs(DOCS_DIR, exist_ok=True)
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", ".", "?", "!", " ", ""]
    )

    all_chunks = []
    supported_extensions = ('.pdf', '.docx', '.png', '.md')
    # 특정 공공기관 CSV 파일은 여기서 제외하고 별도로 로드 처리함
    all_files = [f for f in os.listdir(DOCS_DIR) if f.lower().endswith(supported_extensions) and not f.startswith("보건복지부")]

    if not all_files:
        print(f"[Information] {DOCS_DIR} 폴더에 처리할 지원 문서가 없습니다.")
        return []

    print(f"■ 총 {len(all_files)}개의 문서 분석을 시작합니다.")

    for file_name in all_files:
        file_path = os.path.join(DOCS_DIR, file_name)
        ext = os.path.splitext(file_name)[1].lower()
        
        text = ""
        if ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif ext == '.docx':
            text = extract_text_from_docx(file_path)
        elif ext == '.png':
            text = extract_text_from_png(file_path)
        elif ext == '.md':
            text = extract_text_from_md(file_path)
            
        if not text.strip():
            print(f"  └─ [패스] 추출된 텍스트가 비어있습니다.")
            continue

        docs = splitter.create_documents(
            texts=[text],
            metadatas=[{
                "source": file_name,
                "category": "clinical",
                "char_length": len(text)
            }]
        )
        all_chunks.extend(docs)
        print(f"  └─ 완료: {len(docs)}개의 청크 생성 완료.")

    return all_chunks


# 지역 단축 표준 매핑 딕셔너리
REGION_MAPPING = {
    "서울": "서울", "서울특별시": "서울",
    "경기": "경기", "경기도": "경기",
    "인천": "인천", "인천광역시": "인천",
    "부산": "부산", "부산광역시": "부산",
    "대구": "대구", "대구광역시": "대구",
    "광주": "광주", "광주광역시": "광주",
    "대전": "대전", "대전광역시": "대전",
    "울산": "울산", "울산광역시": "울산",
    "세종": "세종", "세종특별자치시": "세종",
    "강원": "강원", "강원도": "강원", "강원특별자치도": "강원",
    "충북": "충북", "충청북도": "충북",
    "충남": "충남", "충청남도": "충남",
    "전북": "전북", "전라북도": "전북", "전북특별자치도": "전북",
    "전남": "전남", "전라남도": "전남",
    "경북": "경북", "경상북도": "경북",
    "경남": "경남", "경상남도": "경남",
    "제주": "제주", "제주특별자치도": "제주", "제주도": "제주"
}

def parse_region(address: str) -> str:
    """주소에서 지역을 파싱하여 표준 한글 지역명으로 단축 변환합니다."""
    if not address or not isinstance(address, str):
        return "서울"
    
    # 공백으로 잘라 첫 어절을 기준으로 지역명 확인
    tokens = address.strip().split()
    if not tokens:
        return "서울"
        
    first_token = tokens[0]
    
    # REGION_MAPPING에서 정확히 매칭되는지 확인
    for raw_name, std_name in REGION_MAPPING.items():
        if first_token.startswith(raw_name) or raw_name in first_token:
            return std_name
            
    # 매핑되지 않은 경우 기본값 반환
    return "서울"

def load_public_resource_csv() -> List[Document]:
    """공공기관 정신건강 관련기관 CSV 파일을 로드하여 Document 리스트로 파싱합니다."""
    csv_file = None
    for f in os.listdir(DOCS_DIR):
        if f.startswith("보건복지부") and f.endswith(".csv"):
            csv_file = os.path.join(DOCS_DIR, f)
            break
            
    if not csv_file:
        print("[Warning] 보건복지부 정신건강 관련기관 CSV 파일을 찾을 수 없습니다.")
        return []
        
    print(f"■ 공공기관 CSV 분석 시작: {os.path.basename(csv_file)}")
    documents = []
    
    try:
        # CP949 인코딩으로 파일 열기
        with open(csv_file, mode="r", encoding="cp949") as f:
            reader = csv.DictReader(f)
            count = 0
            for row in reader:
                # CSV 헤더 예시: 기관명, 기관구분, 주소, 홈페이지
                name = row.get("기관명", "").strip()
                category = row.get("기관구분", "").strip()
                address = row.get("주소", "").strip()
                homepage = row.get("홈페이지", "").strip()
                
                if not name or not address:
                    continue
                    
                region = parse_region(address)
                
                # 가독성 높은 page_content 구성
                content = (
                    f"기관명: {name}\n"
                    f"기관구분: {category}\n"
                    f"주소: {address}\n"
                    f"홈페이지: {homepage if homepage else '정보 없음'}"
                )
                
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": os.path.basename(csv_file),
                        "category": "public",
                        "region": region,
                        "name": name,
                        "address": address,
                        "homepage": homepage
                    }
                )
                documents.append(doc)
                count += 1
                
        print(f"  └─ 완료: CSV 데이터 {count}건 파싱 및 지역 태깅 완료.")
        return documents
    except Exception as e:
        print(f"[Error] CSV 파싱 실패. 상세: {e}")
        return []

def build_kb():
    """RAG 지식 데이터베이스 구축 파이프라인을 실행합니다."""
    print("===================================================")
    print("   [말랑해도 돼] RAG 상담 지식 DB 구축 파이프라인")
    print("===================================================\n")

    engine = RAGEngine()

    # 1. 문서들을 읽어와 청킹 및 clinical_kb 구축
    chunks = process_and_chunk_documents()
    if chunks:
        print(f"\n■ 총 {len(chunks)}개의 상담 지식 조각(Chunk)을 FAISS 벡터 인덱스(clinical_kb)로 구축합니다...")
        try:
            engine.build_index(chunks, "clinical_kb")
            print(f"🎉 clinical_kb 구축 성공! 경로: {os.path.join(VECTOR_STORE_DIR, 'clinical_kb')}")
        except Exception as e:
            print(f"❌ clinical_kb 인덱스 빌드 실패: {e}")
    else:
        print("\n[알림] 처리할 임상 문서 청크가 없습니다.")

    # 2. 공공기관 CSV 데이터를 읽어와 public_resource_kb 구축
    public_docs = load_public_resource_csv()
    if public_docs:
        print(f"\n■ 총 {len(public_docs)}개의 공공 자원 데이터를 FAISS 벡터 인덱스(public_resource_kb)로 구축합니다...")
        try:
            engine.build_index(public_docs, "public_resource_kb")
            print(f"🎉 public_resource_kb 구축 성공! 경로: {os.path.join(VECTOR_STORE_DIR, 'public_resource_kb')}")
        except Exception as e:
            print(f"❌ public_resource_kb 인덱스 빌드 실패: {e}")
    else:
        print("\n[알림] 처리할 공공기관 자원 데이터가 없습니다.")

    print("\n===================================================")
    print("🎉 RAG 상담 지식 DB 다중 구축 파이프라인 완료!")
    print("===================================================")

if __name__ == "__main__":
    build_kb()
